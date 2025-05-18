
from shapely.geometry import Point
from typing import Dict
import pandas as pd
import math
import random
import numpy as np
import os
import matplotlib
matplotlib.use('Agg')  
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from collections import Counter
from docx import Document
from typing import Tuple
from docx.shared import Inches
from datetime import datetime
from matplotlib.patches import Patch
from collections import defaultdict
import subprocess
import geopandas as gpd
from io import StringIO
import sys
from openpyxl import load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
import openpyxl

#################################################################################################################################
#
#          
#
#
#                     SIMULATEUR ÉLECTORAL DU PÉROU CONFORME À LA LOE (LEY ORGÁNICA DE ELECCIONES)
#
#
#
#
#################################################################################################################################

#################################################################################################################################
# PARAMÈTRES
#################################################################################################################################
SEUIL_ELECTORAL = 5
NB_DEPUTES = 130
NB_SENATEURS_TOTAL = 60
NB_SENATEURS_MAJ = 30
NB_SENATEURS_PROP = 34
NB_SENATEURS_LIMA = 4
NB_SENATEURS_NATIONAL = 30

# Répartition réelle des sièges par district (selon le tableau officiel 2026)
DISTRIBUTION_DEPUTES = {
    "Amazonas": 2, "Ancash": 5, "ApurÃ­mac": 2, "Arequipa": 6, "Ayacucho": 3,
    "Cajamarca": 6, "Callao": 4, "Cusco": 5, "Huancavelica": 2, "HuÃ¡nuco": 3,
    "Ica": 4, "JunÃ­n": 5, "La Libertad": 7, "Lambayeque": 5, "Lima Metropolitana": 32,
    "Lima Provincias": 4, "Exteriores": 2, "Loreto": 4, "Madre de Dios": 2,
    "Moquegua": 2, "Pasco": 2, "Piura": 7, "Puno": 5, "San MartÃ­n": 4,
    "Tacna": 2, "Tumbes": 2, "Ucayali": 3
}

#################################################################################################################################
# MONTE CARLO SIMULATIONS (Députés uniquement)
#################################################################################################################################

def simulation_monte_carlo(pourcentages: Dict[str, float], n_iterations: int = 250) -> pd.DataFrame:
    resultats = []
    pourcentages_filtrés = {p: v for p, v in pourcentages.items() if v >= SEUIL_ELECTORAL}
    for _ in range(n_iterations):
        try:
            perturbés = {p: max(0, v + random.uniform(-2, 2)) for p, v in pourcentages_filtrés.items()}
            total = sum(perturbés.values())
            if total == 0:
                continue
            normalisés = {p: v / total * 100 for p, v in perturbés.items()}
            df = appliquer_cifra_par_district(normalisés, DISTRIBUTION_DEPUTES)
            if not df.empty:
                resultats.append(df.set_index("Parti")["Sièges"])
        except Exception as e:
            print(f"Erreur : {e}")

    if not resultats:
        print("❌ Aucune simulation n’a pu être effectuée.")
        return pd.DataFrame()

    df_resultats = pd.DataFrame(resultats).fillna(0)
    stats = df_resultats.describe(percentiles=[0.05, 0.5, 0.95]).T
    stats.rename(columns={"5%": "Sièges_min", "50%": "Sièges_med", "95%": "Sièges_max"}, inplace=True)
    stats = stats[["Sièges_min", "Sièges_med", "Sièges_max"]].astype(int)
    print("\n============== Simulation Monte Carlo (250 itérations) ==============")
    print(stats.to_string())
    return stats

#################################################################################################################################
# AFFICHE LES LIMITES DU GRAPHIQUE MONTE CARLO
#################################################################################################################################

def afficher_bornes_monte_carlo(stats: pd.DataFrame):
    stats = stats.sort_values("Sièges_med", ascending=True)
    fig, ax = plt.subplots(figsize=(10, 6))

    bars_max = ax.barh(stats.index, stats["Sièges_max"], color="#de4421", label="Max")
    bars_med = ax.barh(stats.index, stats["Sièges_med"], color="#f5e50c", label="Médiane")
    bars_min = ax.barh(stats.index, stats["Sièges_min"], color="#bbf50c", label="Min")

    # Ajouter les étiquettes au bout des barres
    for bars, colonne in zip([bars_max, bars_med, bars_min], ["Sièges_max", "Sièges_med", "Sièges_min"]):
        for bar in bars:
            largeur = bar.get_width()
            ax.text(largeur + 1, bar.get_y() + bar.get_height() / 2,
                    f"{int(largeur)}", va='center', fontsize=9)

    ax.set_xlabel("Nombre de sièges")
    # Ajouter la date et heure en haut à droite
    now = datetime.now().strftime("%d/%m/%Y, %H:%M")
    fig.text(0.985, 0.98, f"{now}", fontsize=10, color='gray', ha='right', va='top')
    
    # Ajouter la mention de copyright en bas à gauche
    fig.text(0.66, 0.01, "© Matthieu Dangeville @matthieu_d0", fontsize=12, color='gray', ha='left', va='bottom')
    
    ax.set_title("Incertitude des projections (250 simulations)", fontsize=14, fontweight="bold")
    ax.legend()
    plt.tight_layout()

    chemin = r"C:\Matthieu\Documents\montecarlo_barres.png"
    plt.savefig(chemin, dpi=300)
    plt.close()

    if os.path.exists(chemin):
        os.startfile(chemin)

#################################################################################################################################
# FONCTION CIFRA REPARTIDORA PAR DISTRICT
#################################################################################################################################

def appliquer_cifra_par_district(
    pourcentages: Dict[str, float],
    distribution: Dict[str, int],
    perturbations: Dict[str, Dict[str, float]] = None  # 🔹 paramètre optionnel
) -> pd.DataFrame:
    result_total = pd.DataFrame(columns=["Parti", "Sièges"])

    for district, nb_sieges in distribution.items():
        # 🔸 Si une perturbation a été fournie pour ce district, on l’utilise
        if perturbations and district in perturbations:
            perturbation = perturbations[district]
        else:
            # 🔸 Sinon, on en génère une nouvelle (par défaut, pour compatibilité)
            perturbation = {p: max(0, v + random.uniform(-2, 2)) for p, v in pourcentages.items()}

        df = pd.DataFrame(perturbation.items(), columns=["Parti", "Pourcentage"])
        total_valid = df["Pourcentage"].sum()
        df["Votes"] = df["Pourcentage"] / total_valid * 1_000_000

        quotients = []
        for idx, row in df.iterrows():
            for i in range(1, nb_sieges + 1):
                quotients.append({
                    "Parti": row["Parti"],
                    "Quotient": row["Votes"] / i
                })

        top = sorted(quotients, key=lambda x: x["Quotient"], reverse=True)[:nb_sieges]
        repartition = pd.DataFrame(top).groupby("Parti").size().reset_index(name="Sièges")
        result_total = pd.concat([result_total, repartition], ignore_index=True)

    return result_total.groupby("Parti").sum().reset_index().sort_values("Sièges", ascending=False)

#################################################################################################################################
# SÉNAT : MAJORITAIRE (1/région sauf Lima) + LIMA à la proportionnelle + NATIONAL à la proportionnelle
#################################################################################################################################

def simulateur_senat(pourcentages_nationaux: Dict[str, float], pourcentages_filtrés: Dict[str, float]) -> pd.DataFrame:
        # Majoritaire : sans filtrage
    regions_maj = [r for r in DISTRIBUTION_DEPUTES if r != "Lima Metropolitana"]
    majoritaire_result = []
    
    for region in regions_maj[:NB_SENATEURS_MAJ]:
        # Perturbation locale des pourcentages (simulation d’un vote régional)
        perturbation = {p: max(0, v + random.uniform(-2, 2)) for p, v in pourcentages_nationaux.items()}
        total = sum(perturbation.values())
        if total == 0:
            continue
        local_votes = {p: (v / total) * 1_000_000 for p, v in perturbation.items()}
        gagnant = max(local_votes.items(), key=lambda x: x[1])[0]
        majoritaire_result.append(gagnant)

    # Compilation des vainqueurs majoritaires
    df_majoritaire = pd.DataFrame(Counter(majoritaire_result).items(), columns=["Parti", "Sièges_majoritaire"])

    # ✅ Dictionnaire des vainqueurs par région (pour carte)
    region_senateurs = {}
    for i, region in enumerate(regions_maj[:NB_SENATEURS_MAJ]):
        if i < len(majoritaire_result):
            region_senateurs[region] = majoritaire_result[i]


    # Proportionnelle : avec filtrage
    df_lima = appliquer_cifra_par_district(pourcentages_filtrés, {"Lima Metropolitana": NB_SENATEURS_LIMA})
    df_lima.rename(columns={"Sièges": "Sièges_Lima"}, inplace=True)

    df_national = appliquer_cifra_par_district(pourcentages_filtrés, {"National": NB_SENATEURS_NATIONAL})
    df_national.rename(columns={"Sièges": "Sièges_National"}, inplace=True)

    # Fusion des trois
    df = pd.merge(df_majoritaire, df_lima, on="Parti", how="outer")
    df = pd.merge(df, df_national, on="Parti", how="outer").fillna(0)
    df["Total_Sénat"] = df["Sièges_majoritaire"] + df["Sièges_National"] + df["Sièges_Lima"]    
    df["Total_Sénat"] = df["Total_Sénat"].astype(int)
    
    # ⚠️ Appliquer le double seuil : 5 % des voix (déjà fait) + au moins 3 sièges au total au Sénat
    df = df[df["Total_Sénat"] >= 3]

    # On récupère les vainqueurs spécifiques
    vainqueur_exteriores = None
    if "Exteriores" in DISTRIBUTION_DEPUTES:
        perturbation_ext = {p: max(0, v + random.uniform(-2, 2)) for p, v in pourcentages_nationaux.items()}
        total_ext = sum(perturbation_ext.values())
        if total_ext > 0:
            local_votes_ext = {p: (v / total_ext) * 1_000_000 for p, v in perturbation_ext.items()}
            vainqueur_exteriores = max(local_votes_ext.items(), key=lambda x: x[1])[0]
            region_senateurs["Exteriores"] = vainqueur_exteriores
    
    # Pour Lima Metropolitana (4 sièges à la proportionnelle)
    if not df_lima.empty:
        vainqueur_lima = df_lima.sort_values("Sièges_Lima", ascending=False).iloc[0]["Parti"]
        region_senateurs["Lima Metropolitana"] = vainqueur_lima
    else:
        print("❌ Aucune donnée valide pour Lima Metropolitana")

    return df.sort_values("Total_Sénat", ascending=False), region_senateurs

#################################################################################################################################
# CARTE DE L'ELECTION DU SENAT - PARTIE MAJORITAIRE
#################################################################################################################################

def generer_carte_senateurs_majoritaires(gagnants_region: Dict[str, str], fichier_sortie: str):
    gdf = gpd.read_file("PER_adm1.shp")
    gdf["region_nom"] = gdf["NAME_1"]
    mapping_noms = {
        "Lima": "Lima Metropolitana",
        "El Callao": "Callao",
        "La Libertad": "La Libertad",
        "San Martín": "San Martin",
        "Lima Province": "Lima Provincias"
    }
    gdf["region_match"] = gdf["region_nom"].replace(mapping_noms)
    gdf["centroid"] = gdf.geometry.centroid

    couleurs = {
        "Perú Libre": "#d4120f",
        "Fuerza Popular": "#e67710",
        "Acción Popular": "#000000",
        "Alianza para el Progreso": "#1b49b5",
        "Renovación Popular": "#49b9de",
        "Avanza Pais": "#d62b95",
        "PPC": "#42a912",
        "PNP": "#e75b42",
        "Juntos por el Perú": "#3fd62b",
        "FA": "#336c0c",
        "Somos Perú": "#2b42d6",
        "Podemos Perú": "#dba321",
        "Partido Morado": "#bf21db",
        "APRA": "#e8e11e",
        "Victoria Nacional": "#ec3710",
        "Frepap": "#102a9f",
        "Unión por el Perú": "#e12d21"
    }

    geometries = []
    couleurs_points = []
    partis_points = []

    for _, row in gdf.iterrows():
        region = row["region_match"]
    
        if region == "Lima Metropolitana":
            print("")
        else:
            if region in gagnants_region:
                centroid = row["centroid"]
                parti = gagnants_region[region]
                if abs(centroid.x - (-76.30)) < 0.01 and abs(centroid.y - (-12.02)) < 0.01:
                    print("")
    
                geometries.append(centroid)
                couleurs_points.append(couleurs.get(parti, "#808080"))
                partis_points.append(parti)

        # --- Exteriores (1 siège fictif en haut gauche) ---
    if "Exteriores" in gagnants_region:
        geometries.append(Point(-80, -2.5))
        couleurs_points.append(couleurs.get(gagnants_region["Exteriores"], "#808080"))
        partis_points.append(gagnants_region["Exteriores"])

    if "Lima Metropolitana" in gagnants_region:
        parti_lima = gagnants_region["Lima Metropolitana"]
        
    if "Lima Provincias" in gagnants_region:
        print("")

    # --- Lima Metropolitana (4 sièges proportionnels) ---
    lima_row = gdf[gdf["region_match"] == "Lima Metropolitana"]
    if not lima_row.empty:
        if "Lima Metropolitana" in gagnants_region:
            parti_lima = gagnants_region["Lima Metropolitana"]
            couleur_lima = couleurs.get(parti_lima, "#808080")
            lima_centroid = lima_row.iloc[0]["centroid"]  # ✅ cette ligne était manquante !
            for i in range(4):
                x = lima_centroid.x - 3 + i * 0.6
                y = lima_centroid.y - 0.3
                geometries.append(Point(x, y))
                couleurs_points.append(couleur_lima)
                partis_points.append(parti_lima)
    
    # --- Lima Provincias (1 siège majoritaire) ---
    if "Lima Provincias" in gagnants_region:
        prov_row = gdf[gdf["region_match"] == "Lima Provincias"]
        if not prov_row.empty:
            centroid = prov_row.iloc[0]["centroid"]
            centroid = Point(centroid.x + 0.6, centroid.y + 0)
            geometries.append(centroid)
            couleur_prov = couleurs.get(gagnants_region["Lima Provincias"], "#808080")
            couleurs_points.append(couleur_prov)
            partis_points.append(gagnants_region["Lima Provincias"])


    df_points = gpd.GeoDataFrame(geometry=geometries)
    df_points["Couleur"] = couleurs_points
    df_points["Parti"] = partis_points

        # 🔴 Suppression du point parasite (valeurs à adapter si besoin)
    coord_x_suspecte = -76.90
    coord_y_suspecte = -12.02
    print("Point suspect supprimé. Chargement Monte Carlo ...")
    tolérance = 0.01  # tolérance autour des coordonnées (peut être ajustée)
    
    # Création d’un masque pour détecter le point suspect
    masque_valide = ~df_points.geometry.apply(
        lambda p: abs(p.x - coord_x_suspecte) < tolérance and abs(p.y - coord_y_suspecte) < tolérance
    )
    
    # Filtrage des points
    df_points = df_points[masque_valide].reset_index(drop=True)
    partis_points = df_points["Parti"].tolist()
    couleurs_points = df_points["Couleur"].tolist()

    fig, ax = plt.subplots(figsize=(12, 12))
    gdf.boundary.plot(ax=ax, linewidth=0.8, color="black")
    df_points.plot(ax=ax, color=df_points["Couleur"], markersize=200)
    
    # Calculer le nombre de points (donc sénateurs) par parti
    compteur_senateurs = Counter(partis_points)
    
    # Construire une nouvelle légende avec le total de sièges
    legendes = [
        Patch(color=couleurs.get(p, "#808080"), label=f"{p} ({compteur_senateurs.get(p, 0)})")
        for p in sorted(compteur_senateurs.keys(), key=lambda x: -compteur_senateurs[x])
    ]
    
    plt.legend(handles=legendes, loc="lower left", fontsize=9, title="Sénateurs élus")
    # Ajouter la date et heure en haut à droite
    now = datetime.now().strftime("%d/%m/%Y, %H:%M")
    fig.text(0.985, 0.98, f"{now}", fontsize=10, color='gray', ha='right', va='top')
    
    # Ajouter la mention de copyright en bas à gauche
    fig.text(0.65, 0.01, "© Matthieu Dangeville @matthieu_d0", fontsize=15, color='gray', ha='left', va='bottom')
    
    plt.title("Sénat par circonscription majoritaire sauf Lima proportionnelle (30)", fontsize=14, fontweight="bold")
    plt.axis("off")
    plt.tight_layout()
    
    if os.path.exists(fichier_sortie):
        os.remove(fichier_sortie)  # Supprime l'image précédente    
    plt.savefig(fichier_sortie, dpi=300)
    plt.close()

    if os.path.exists(fichier_sortie):
        try:
            os.startfile(fichier_sortie)
        except AttributeError:
            import subprocess
            subprocess.run(["open", fichier_sortie])  # pour macOS


#################################################################################################################################
# COMPARE LE RESULTAT DES DEPUTES AVEC 2021
#################################################################################################################################

def comparer_resultats(simules_df: pd.DataFrame, historiques: Dict[str, int]):
    comparatif = []
    simules_dict = dict(zip(simules_df["Parti"], simules_df["Sièges"]))
    for parti in set(historiques) | set(simules_dict):
        s2021 = historiques.get(parti, 0)
        s_sim = simules_dict.get(parti, 0)
        difference = s_sim - s2021
        comparatif.append((parti, s2021, s_sim, difference))
    return sorted(comparatif, key=lambda x: x[2], reverse=True)

#################################################################################################################################
# SIMULATION 1ER TOUR ELECTION PRESIDENTIELLE
#################################################################################################################################

def simuler_premier_tour(pourcentages: Dict[str, float]):
    valides = {p: v for p, v in pourcentages.items() if v > 0}
    total = sum(valides.values())
    normalisés = {p: v / total * 100 for p, v in valides.items()}
    classement = sorted(normalisés.items(), key=lambda x: x[1], reverse=True)
    return classement[:2], normalisés  # Renvoie les deux finalistes

#################################################################################################################################
# SIMULATION 2E TOUR ELECTION PRESIDENTIELLE
#################################################################################################################################

def simuler_second_tour(finalistes: Tuple[str, str], pourcentages: Dict[str, float]):
    cand1, cand2 = finalistes
    autres = [p for p in pourcentages if p not in finalistes]

    # Hypothèse simple : répartition équitable des reports
    reports_cand1 = sum(pourcentages[a] * 0.5 for a in autres)
    reports_cand2 = sum(pourcentages[a] * 0.5 for a in autres)

    score1 = pourcentages[cand1] + reports_cand1
    score2 = pourcentages[cand2] + reports_cand2

    total = score1 + score2
    pct1 = score1 / total * 100
    pct2 = score2 / total * 100

    return {
        cand1: round(pct1, 2),
        cand2: round(pct2, 2),
        "gagnant": cand1 if pct1 > pct2 else cand2
    }

#################################################################################################################################
# LANCER LA SIMULATION DE L'ELECTION PRESIDENTIELLE
#################################################################################################################################

def simuler_election_presidentielle(pourcentages: Dict[str, float]):
    print("============== 🗳️ ÉLECTION PRÉSIDENTIELLE ==============")

    # 🔹 Perturbation aléatoire de ±3 % directement ici
    perturbés = {p: max(0, v + random.uniform(-3, 3)) for p, v in pourcentages.items()}
    total = sum(perturbés.values())
    perturbés = {p: v / total * 100 for p, v in perturbés.items()}  # normalisation à 100 %

    # 🔹 Premier tour avec scores perturbés
    finalistes, normalisés = simuler_premier_tour(perturbés)

    print("\n📊 Résultats du 1er tour :")
    classement_complet = sorted(normalisés.items(), key=lambda x: x[1], reverse=True)
    for parti, pct in classement_complet:
        print(f" - {parti} : {pct:.2f} %")

    second_tour = simuler_second_tour((finalistes[0][0], finalistes[1][0]), normalisés)

    print("\n⚔️  Second tour :")
    print(f"{finalistes[0][0]} : {second_tour[finalistes[0][0]]:.2f}%")
    print(f"{finalistes[1][0]} : {second_tour[finalistes[1][0]]:.2f}%")
    print(f"\n🏆 Gagnant : {second_tour['gagnant']}")

#################################################################################################################################
# GENERE UN GRAPHIQUE A BARRES DES SENATEURS ELUS - PARTIE PROPORTIONNELLE
#################################################################################################################################

def graphique_senat_proportionnel(df_senat, chemin_image="C:/Matthieu/Documents/senat_proportionnel_barres.png"):
    import matplotlib.pyplot as plt

    couleurs = {
        "Perú Libre": "#d4120f",
        "Fuerza Popular": "#e67710",
        "Acción Popular": "#000000",
        "Alianza para el Progreso": "#1b49b5",
        "Renovación Popular": "#49b9de",
        "Avanza Pais": "#d62b95",
        "PPC": "#42a912",
        "PNP": "#e75b42",
        "Juntos por el Perú": "#3fd62b",
        "FA": "#336c0c",
        "Somos Perú": "#2b42d6",
        "Podemos Perú": "#dba321",
        "Partido Morado": "#bf21db",
        "APRA": "#e8e11e",
        "Victoria Nacional": "#ec3710",
        "Frepap": "#102a9f",
        "Unión por el Perú": "#e12d21"
    }
    # Extraire uniquement les colonnes proportionnelles
    df = df_senat.copy()
    df["Sièges_Prop"] = df["Sièges_National"]
    df = df[df["Sièges_Prop"] > 0].sort_values("Sièges_Prop", ascending=True)

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.barh(df["Parti"], df["Sièges_Prop"], color=[couleurs.get(p, "#808080") for p in df["Parti"]])

    # Ajouter les étiquettes de valeur
    for bar in bars:
        width = bar.get_width()
        ax.text(width + 0.3, bar.get_y() + bar.get_height()/2, f"{int(width)}", va='center', fontsize=10)

    ax.set_title("Sénat – Proportionnelle nationale (30)", fontsize=14, fontweight="bold")
    ax.set_xlabel("Nombre de sièges")
    
    # Ajouter la date et heure en haut à droite
    now = datetime.now().strftime("%d/%m/%Y, %H:%M")
    fig.text(0.985, 0.98, f"{now}", fontsize=10, color='gray', ha='right', va='top')
    
    # Ajouter la mention de copyright en bas à gauche
    fig.text(0.65, 0.01, "© Matthieu Dangeville @matthieu_d0", fontsize=12, color='gray', ha='left', va='bottom')
    
    plt.tight_layout()
    plt.savefig(chemin_image, dpi=300)
    plt.close()

    if os.path.exists(chemin_image):
        try:
            os.startfile(chemin_image)
        except AttributeError:
            import subprocess
            subprocess.run(["open", chemin_image])

#################################################################################################################################
# VISUALISATION : DEMI-CERCLE DES DÉPUTÉS
#################################################################################################################################

def generate_double_half_circle(df_deputes, df_senat):


# Couleurs par parti
    couleurs = {
        "Perú Libre": "#d4120f",
        "Fuerza Popular": "#e67710",
        "Acción Popular": "#000000",
        "Alianza para el Progreso": "#1b49b5",
        "Renovación Popular": "#49b9de",
        "Avanza Pais": "#d62b95",
        "PPC": "#42a912",
        "PNP": "#e75b42",
        "Juntos por el Perú": "#3fd62b",
        "FA": "#336c0c",
        "Somos Perú": "#2b42d6",
        "Podemos Perú": "#dba321",
        "Partido Morado": "#bf21db",
        "APRA": "#e8e11e",
        "Victoria Nacional": "#ec3710",
        "Frepap": "#102a9f",
        "Unión por el Perú": "#e12d21"
    }



    # Ordre idéologique
    ordre_politique = [
        "Renovación Popular",
        "Fuerza Popular",
        "Avanza Pais",
        "APRA",
        "PPC",
        "Podemos Perú",
        "Alianza para el Progreso",
        "Frepap",
        "Acción Popular",
        "Partido Morado",
        "Somos Perú",
        "FA",
        "Victoria Nacional",
        "Unión por el Perú",
        "Juntos por el Perú",
        "PNP",
        "Perú Libre"
    ]

#################################################################################################################################
# TRACE LES HEMICYCLES
#################################################################################################################################

    def tracer_hemicycle(ax, df, titre):
        df = df[df["Sièges"] > 0].copy()
        df["ordre"] = df["Parti"].apply(lambda p: ordre_politique.index(p) if p in ordre_politique else 999)
        df = df.sort_values("ordre").drop(columns="ordre")
    
        total = df["Sièges"].sum()
        angles = df["Sièges"] / total * 180
    
        start_angle = 0
        handles = []
        labels = []
    
        for row, angle in zip(df.itertuples(index=False), angles):
            theta = np.linspace(np.radians(start_angle), np.radians(start_angle + angle), 100)
            x = np.append([0], np.cos(theta))
            y = np.append([0], np.sin(theta))
            couleur = couleurs.get(row.Parti, "#808080")
            if couleur == "#808080":
                print(f"⚠️ ATTENTION : Parti inconnu détecté : {row.Parti} ⚠️")
    
            patch = ax.fill(x, y, color=couleur, label=f"{row.Parti} ({row.Sièges})")
            handles += patch
            labels.append(f"{row.Parti} ({row.Sièges})")
            start_angle += angle
    
    
        ax.set_aspect('equal')
        ax.axis('off')
        ax.set_title(titre, fontsize=25, fontweight='bold', pad=25)
        return handles, labels
    
    # Préparation
    df_dep = df_deputes.rename(columns={"Sièges": "Sièges"})
    # df_sen = df_senat.rename(columns={"Total_Sénat": "Sièges"})
    
        # Récupération des colonnes nécessaires pour l’hémicycle
    df_sen = df_senat.copy()
    
    df_sen = df_sen[["Parti", "Sièges"]]  # C'est CE dataframe qu'il faut passer
    
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(15, 11))
    
    # Tracer Sénat et Députés avec légendes
    handles_sen, labels_sen = tracer_hemicycle(ax1, df_sen, "Senado (60 escaños)")
    handles_dep, labels_dep = tracer_hemicycle(ax2, df_dep, "Camara de Diputados (130 escaños)")
    
    # Légendes distinctes
    ax1.legend(handles_sen, labels_sen, loc='center left', bbox_to_anchor=(1, 0.4), fontsize=15)
    ax2.legend(handles_dep, labels_dep, loc='center left', bbox_to_anchor=(1, 0.8), fontsize=15)
    
    # Ajouter la date et heure en haut à droite
    now = datetime.now().strftime("%d/%m/%Y, %H:%M")
    fig.text(0.985, 0.98, f"{now}", fontsize=10, color='gray', ha='right', va='top')
    
    # Ajouter la mention de copyright en bas à gauche
    fig.text(0.73, 0.01, "© Matthieu Dangeville @matthieu_d0", fontsize=15, color='gray', ha='left', va='bottom')
    
    # === Nouvelle légende : Écarts par rapport à 2021 ===
    df_dep_dict = dict(zip(df_dep["Parti"], df_dep["Sièges"]))
    liste_ecarts = []
    for parti in df_dep_dict:
        anciens = resultats_2021.get(parti, 0)
        nouveaux = df_dep_dict[parti]
        ecart = nouveaux - anciens
        signe = "+" if ecart > 0 else ""
        texte = f"{parti}: {signe}{ecart}"
        liste_ecarts.append(texte)
    
    # Affichage sous la légende des députés (ajustement possible)
    ax2.text(1.05, 0.4, "Écart par rapport à 2021 :", fontsize=14, fontweight='bold', transform=ax2.transAxes)
    for i, texte in enumerate(liste_ecarts):
        ax2.text(1.05, 0.35 - i*0.05, texte, fontsize=13, transform=ax2.transAxes)
    
    
    plt.tight_layout(h_pad=2.5)
    
    chemin_image = r"C:\Matthieu\Documents\double_hemicycle.png"
    plt.savefig(chemin_image, dpi=300)
    plt.close()
    
    if os.path.exists(chemin_image):
        try:
            os.startfile(chemin_image)
        except AttributeError:
            subprocess.run(["open", chemin_image])
    else:
        print("\n❌ L’image n’a pas été enregistrée.")
        print()


resultats_2021 = {
    "Perú Libre": 37,
    "Fuerza Popular": 24,
    "Renovación Popular": 13,
    "Acción Popular": 16,
    "Alianza para el Progreso": 15,
    "Avanza Pais": 7,
    "Juntos por el Perú": 5,
    "Somos Perú": 5,
    "Podemos Perú": 5,
    "Partido Morado": 3
}

#################################################################################################################################
# GENERE UN FICHIER WORD
#################################################################################################################################

def generer_rapport_word(chemin_dossier: str, texte_console: str,
                         image_hemicycle_path: str,
                         image_montecarlo_path: str,
                         image_carte_path: str,
                         df_deputes: pd.DataFrame,
                         df_senat: pd.DataFrame,
                         stats_monte_carlo: pd.DataFrame,
                         nom_fichier: str = "Rapport_simulation_python_Peru.docx"):


    # Créer un document Word
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)     # marge supérieure réduite à 0.5 pouces
        section.bottom_margin = Inches(0.3)  # marge inférieure réduite à 0.5 pouces

    
    doc.add_heading("Rapport de simulation électorale – Pérou", 0)
    doc.add_paragraph(f"Date de génération : {datetime.now().strftime('%d/%m/%Y à %H:%M')}. Fait par Matthieu Dangeville.")
    doc.add_paragraph("Conforme à la Ley Orgánica de Elecciones (LOE) de la République du Pérou.")
    
    intro = (
        "Ce rapport présente les résultats d’une simulation électorale complète réalisée à l’aide d’un simulateur Python conçu conformément aux principes et exigences de la Ley Orgánica de Elecciones (LOE) du Pérou.\n\n"
        "Le programme modélise de manière rigoureuse l’allocation des sièges au Congrès péruvien, composé de la Chambre des députés (130 membres) et d’un Sénat réinstauré dans sa forme bicamérale (60 membres), tel que prévu dans les réformes institutionnelles en discussion.\n\n" 
        "La répartition des 130 députés par district respecte l'article 21 de la LOE, fondé sur la population électorale recensée. Chaque département, ainsi que la circonscription des électeurs résidant à l’étranger, se voit attribuer un nombre fixe de sièges, comme établi dans la loi électorale. \n\n"
        "L’attribution des sièges au sein de chaque circonscription repose sur la méthode de la plus forte moyenne (méthode de D’Hondt), conformément à l’article 20 de la LOE, en tenant compte d’un seuil électoral national de 5% des votes valides (article 13-A).\n\n"
        "Pour la composition du Sénat, le programme simule un système mixte (moitié majoritaire, moitié proportionnel) : 26 sénateurs sont élus selon un scrutin uninominal majoritaire par circonscription (à raison d’un sénateur par département), tandis que les 34 autres sont élus à la proportionnelle sur deux circonscriptions : Lima métropolitaine (4 sièges) et la liste nationale (30 sièges).\n\n"
        "Ce simulateur intègre également une modélisation des incertitudes par méthode Monte Carlo, qui permet d’évaluer la robustesse des projections à partir d’intentions de vote sujettes à variation.\n\n"
        "Une série de 100 à 1000 itérations introduit une perturbation aléatoire autour des intentions exprimées, simulant l’effet des marges d’erreur, des indécis ou des transferts d’électeurs.\n"
        )
    doc.add_paragraph(intro)


    doc.add_heading("Répartition des députés", level=2)
    table_dep = doc.add_table(rows=1, cols=2)
    table_dep.style = 'Table Grid'
    hdr_cells = table_dep.rows[0].cells
    hdr_cells[0].text = 'Parti'
    hdr_cells[1].text = 'Sièges'
    
    for _, row in df_deputes.iterrows():
        row_cells = table_dep.add_row().cells
        row_cells[0].text = row['Parti']
        row_cells[1].text = str(row['Sièges'])


    doc.add_heading("2. Répartition au Parlement", level=1)
    doc.add_picture(image_hemicycle_path, width=Inches(6.5))


    doc.add_heading("Répartition au Sénat", level=2)
    table_sen = doc.add_table(rows=1, cols=2)
    table_sen.style = 'Table Grid'
    hdr_cells = table_sen.rows[0].cells
    hdr_cells[0].text = 'Parti'
    hdr_cells[1].text = 'Total Sénat'
    
    for _, row in df_senat.iterrows():
        row_cells = table_sen.add_row().cells
        row_cells[0].text = row['Parti']
        row_cells[1].text = str(row['Total_Sénat'])
        
    doc.add_heading("Simulation Monte Carlo", level=2)
    table_mc = doc.add_table(rows=1, cols=4)
    table_mc.style = 'Table Grid'
    hdr_cells = table_mc.rows[0].cells
    hdr_cells[0].text = 'Parti'
    hdr_cells[1].text = 'Sièges min'
    hdr_cells[2].text = 'Sièges médian'
    hdr_cells[3].text = 'Sièges max'

    for parti, row in stats_monte_carlo.iterrows():
        row_cells = table_mc.add_row().cells
        row_cells[0].text = parti
        row_cells[1].text = str(row['Sièges_min'])
        row_cells[2].text = str(row['Sièges_med'])
        row_cells[3].text = str(row['Sièges_max'])

    doc.add_heading("3. Incertitudes Monte Carlo", level=1)
    doc.add_picture(image_montecarlo_path, width=Inches(6.5))

    doc.add_heading("4. Carte des partis dominants par département", level=1)
    doc.add_picture(image_carte_path, width=Inches(6.5))

    doc.add_heading("5. Carte des sénateurs élus par région (majoritaire + Lima)", level=1)
    doc.add_picture("C:/Matthieu/Documents/carte_senateurs.png", width=Inches(6.5))

    doc.add_heading("6. Répartition proportionnelle au Sénat", level=1)
    doc.add_picture("C:/Matthieu/Documents/senat_proportionnel_barres.png", width=Inches(6.5))

    chemin_final = os.path.join(chemin_dossier, nom_fichier)
    doc.save(chemin_final)
    print("\n✅ Rapport Word enregistré. \n FIN DU PROGRAMME.")

#################################################################################################################################
# ENREGISTRE LES RESULTATS DE CHAQUE SIMULATION DANS UNE LIGNE EXCEL !
#################################################################################################################################

def enregistrer_resultats_excel(df_deputes, df_senat, stats_monte_carlo, fichier_excel="C:/Matthieu/Documents/Rapports_des_tests.xlsx"):
    
    date_simulation = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # Création d’un DataFrame consolidé
    df_deputes = df_deputes.set_index("Parti").rename(columns={"Sièges": "Députés"})
    df_senat = df_senat.set_index("Parti").rename(columns={"Total_Sénat": "Sénateurs"})[["Sénateurs"]]
    stats_monte_carlo = stats_monte_carlo.rename(columns={
        "Sièges_min": "MC_min",
        "Sièges_med": "MC_med",
        "Sièges_max": "MC_max"
    })

    df_fusion = df_deputes.join(df_senat, how="outer").join(stats_monte_carlo, how="outer").fillna(0).astype(int)
    df_fusion.insert(0, "Date", date_simulation)

    # Ajout au fichier Excel
    if os.path.exists(fichier_excel):
        wb = load_workbook(fichier_excel)
        if "Simulations" not in wb.sheetnames:
            ws = wb.create_sheet("Simulations")
        else:
            ws = wb["Simulations"]
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Simulations"

    # Si le fichier est vide, écrire l’en-tête
    if ws.max_row == 1 and ws.max_column == 1 and ws["A1"].value is None:
        ws.append(["Date", "Parti", "Députés", "Sénateurs", "MC_min", "MC_med", "MC_max"])

    for parti, row in df_fusion.iterrows():
        ws.append([
            row["Date"],
            parti,
            row.get("Députés", 0),
            row.get("Sénateurs", 0),
            row.get("MC_min", 0),
            row.get("MC_med", 0),
            row.get("MC_max", 0)
        ])
    # ➕ Ajouter une ligne de séparation visuelle
    ws.append(["---- FIN SIMULATION ----", "", "", "", "", "", ""])
    wb.save(fichier_excel)
    print(f"✅ Résultats enregistrés dans {fichier_excel}")


#################################################################################################################################
# GENERE UNE CARTE DES DEPUTES ELUS
#################################################################################################################################

def generer_carte_par_sieges(resultats_par_region: Dict[str, Dict[str, int]], fichier_sortie: str):

    offsets_regionaux = {
    "Lima Metropolitana": (-0, 0),
    "Callao": (-0.3, 0),           # légèrement à gauche
    "Lima Provincias": (0.6, 0.1),   # un peu plus à droite
}
    partis_points = []  # ← Nouveau

    couleurs = {
        "Perú Libre": "#d4120f",
        "Fuerza Popular": "#e67710",
        "Acción Popular": "#000000",
        "Alianza para el Progreso": "#1b49b5",
        "Renovación Popular": "#49b9de",
        "Avanza Pais": "#d62b95",
        "PPC": "#42a912",
        "PNP": "#e75b42",
        "Juntos por el Perú": "#3fd62b",
        "FA": "#336c0c",
        "Somos Perú": "#2b42d6",
        "Podemos Perú": "#dba321",
        "Partido Morado": "#bf21db",
        "APRA": "#e8e11e",
        "Victoria Nacional": "#ec3710",
        "Frepap": "#102a9f",
        "Unión por el Perú": "#e12d21"
    }

    gdf = gpd.read_file("PER_adm1.shp")
    gdf["region_nom"] = gdf["NAME_1"]
    mapping_noms = {
        "Lima": "Lima Metropolitana",
        "El Callao": "Callao",
        "La Libertad": "La Libertad",
        "San Martín": "San Martin",
        "Lima Province": "Lima Provincias"
    }
    gdf["region_match"] = gdf["region_nom"].replace(mapping_noms)
    gdf["centroid"] = gdf.geometry.centroid

    geometries = []
    couleurs_points = []

    for _, row in gdf.iterrows():
        region = row["region_match"]
        centroid = row["centroid"]
        if region not in resultats_par_region:
            continue

        parti_seats = resultats_par_region[region]
        total_seats = sum(parti_seats.values())

        # Taille de la grille : carré le plus petit possible
        n_cols = math.ceil(math.sqrt(total_seats))
        n_rows = math.ceil(total_seats / n_cols)

        spacing = 0.3  # écart entre les points
        offset_x = -(n_cols - 1) / 2 * spacing
        offset_y = -(n_rows - 1) / 2 * spacing

        seat_index = 0
        for parti, count in parti_seats.items():
            for _ in range(count):
                row_idx = seat_index // n_cols
                col_idx = seat_index % n_cols
                dx, dy = offsets_regionaux.get(region, (0, 0))
                x = centroid.x + offset_x + col_idx * spacing + dx
                y = centroid.y + offset_y + row_idx * spacing + dy

                # Décalage manuel pour Lima
                if region == "Lima Metropolitana":
                    x -= 2.5  # Ajuste cette valeur si nécessaire

                geometries.append(Point(x, y))
                couleurs_points.append(couleurs.get(parti, "#808080"))
                partis_points.append(parti)  # ← Nouveau
                seat_index += 1

            # === Ajout manuel de la circonscription "Exteriores" ===
        # Coordonnées fictives en haut à gauche
        x_base, y_base = -80, -2.5  # Ajuste si nécessaire
        
        # Détermine le nombre de sièges pour chaque parti "exteriores"
        sieges_exteriores = resultats_par_region.get("Exteriores", {})
        dx = 0.4  # espacement horizontal entre les cercles
        
        for i, (parti, n_sieges) in enumerate(sieges_exteriores.items()):
            for j in range(n_sieges):
                geometries.append(Point(x_base + j * dx + i * 1.2, y_base))
                couleurs_points.append(couleurs.get(parti, "#808080"))
                partis_points.append(parti)

    df_points = gpd.GeoDataFrame(geometry=geometries)
    df_points["Couleur"] = couleurs_points
    df_points["Parti"] = partis_points
    
    # ✅ Vérification des partis non colorés (gris #808080)
    partis_non_repertories = set()
    for _, row in df_points.iterrows():
        if row["Couleur"] == "#808080":
            partis_non_repertories.add(row["Parti"])

    if partis_non_repertories:
        print("⚠️ Partis sans couleur définie :")
        for p in partis_non_repertories:
            print(f" - {p}")

    fig, ax = plt.subplots(figsize=(12, 12))
    gdf.boundary.plot(ax=ax, linewidth=0.8, color="black")
    df_points.plot(ax=ax, color=df_points["Couleur"], markersize=90)

    plt.title("Répartition des sièges par région (députés)", fontsize=14, fontweight="bold")
    plt.axis("off")

    # Compter les sièges par parti
    # Reconstruire les vrais totaux depuis resultats_par_region
    sieges_totaux = Counter()
    for district, resultats in resultats_par_region.items():
        for parti, nb in resultats.items():
            sieges_totaux[parti] += nb

# Créer la légende avec le nombre de sièges
    legendes = [
    mpatches.Patch(color=c, label=f"{p} ({sieges_totaux.get(p, 0)})")
    for p, c in couleurs.items() if p in sieges_totaux
]

    # Ajouter la date et heure en haut à droite
    now = datetime.now().strftime("%d/%m/%Y, %H:%M")
    fig.text(0.985, 0.98, f"{now}", fontsize=10, color='gray', ha='right', va='top')
    
    # Ajouter la mention de copyright en bas à gauche
    fig.text(0.65, 0.01, "© Matthieu Dangeville @matthieu_d0", fontsize=15, color='gray', ha='left', va='bottom')
    
    
    plt.legend(handles=legendes, loc="lower left", fontsize=9, title="Partis")
    plt.tight_layout()
    plt.savefig(fichier_sortie, dpi=300)
    plt.close()

    print(f"✅ Carte par siège enregistrée : {fichier_sortie}")

    if os.path.exists(fichier_sortie):
        try:
            os.startfile(fichier_sortie)
        except AttributeError:
            import subprocess
            subprocess.run(["open", fichier_sortie])  # pour macOS

#################################################################################################################################
# PONDERATIONS
#################################################################################################################################

"""
# -------------------------------
# PONDÉRATIONS MANUELLES GLOBALES
# -------------------------------
pondérations_personnelles = {
    "Perú Libre": 1.15,
    "Fuerza Popular": 0.85,
    "Acción Popular": 0.95,
    "Alianza para el Progreso": 1.6,
    "Renovación Popular": 0.8,
    "Avanza Pais": 0.75,
    "Juntos por el Perú": 1.05,
    "Somos Perú": 1.1,
    "Podemos Perú": 1.2,
    "Partido Morado": 1.1,
    "APRA": 1.0
}

# -------------------------------
# SIMULATION GLOBALE AVEC GRAPHIQUE
# -------------------------------


# === 🔵 PONDÉRATIONS MANUELLES ===
    pondérations_personnelles = {
    "Perú Libre": 1.15,
    "Fuerza Popular": 0.85,
    "Acción Popular": 0.95,
    "Alianza para el Progreso": 1.6,
    "Renovación Popular": 0.8,
    "Avanza Pais": 0.75,
    "Juntos por el Perú": 1.05,
    "Somos Perú": 1.1,
    "Podemos Perú": 1.2,
    "Partido Morado": 1.1,
    "APRA": 1.0
}


    
    # === 🔵 Application pondérations manuelles sur les députés ===
    deputes["Sièges"] = deputes.apply(
        lambda row: int(round(row["Sièges"] * pondérations_personnelles.get(row["Parti"], 1.0))), axis=1
)


    # === 🔵 Application pondérations manuelles sur le Sénat ===
    senat["Total_Sénat"] = senat.apply(
        lambda row: int(round(row["Total_Sénat"] * pondérations_personnelles.get(row["Parti"], 1.0))), axis=1
)
"""

#################################################################################################################################
# FUSIONNE LES RESULTATS DU SENAT SUR L'HEMICYCLE
#################################################################################################################################

"""
def fusionner_senat(senat_prop_df: pd.DataFrame, gagnants_majoritaires: Dict[str, str]) -> pd.DataFrame:
    
  #  Fusionne les résultats proportionnels (df) et majoritaires (dict) pour affichage dans l’hémicycle du Sénat.
    
    total_majoritaire = Counter(gagnants_majoritaires.values())  # Ex: {"Fuerza Popular": 8, ...}
    total_proportionnel = dict(
        zip(
            senat_prop_df["Parti"],
            senat_prop_df.get("Sièges_National", pd.Series(0)) + senat_prop_df.get("Sièges_Lima", pd.Series(0))
        )
    )

    fusion = {}
    tous_partis = set(total_majoritaire) | set(total_proportionnel)
    for parti in tous_partis:
        fusion[parti] = total_majoritaire.get(parti, 0) + total_proportionnel.get(parti, 0)

    return pd.DataFrame({
        "Parti": list(fusion.keys()),
        "Sièges": list(fusion.values())
    }).sort_values("Sièges", ascending=False)
"""

#################################################################################################################################
# SIMULATION GENERALE
#################################################################################################################################

def simulation_globale(pourcentages: Dict[str, float]):
    # Étape 1 : filtrage initial sur le seuil de 5 %
    candidats_5pourcent = {p: v for p, v in pourcentages.items() if v >= SEUIL_ELECTORAL}
 
    # === 1. Générer les perturbations par district
    perturbations_par_district = {
        district: {p: max(0, v + random.uniform(-2, 2)) for p, v in pourcentages.items()}
        for district in DISTRIBUTION_DEPUTES
    }
    
    # === 2. Répartir les sièges région par région (sans filtrer pour l’instant)
    resultats_par_region = {}
    for district in DISTRIBUTION_DEPUTES:
        df_local = appliquer_cifra_par_district(
            pourcentages,
            {district: DISTRIBUTION_DEPUTES[district]},
            perturbations=perturbations_par_district
        )
        resultats_par_region[district] = dict(zip(df_local["Parti"], df_local["Sièges"]))
    
    # === 3. Agréger tous les résultats régionaux
    totaux_deputes_depuis_carte = defaultdict(int)
    for region, resultats in resultats_par_region.items():
        for parti, nb_sieges in resultats.items():
            totaux_deputes_depuis_carte[parti] += nb_sieges
    
    # === 4. Créer DataFrame des députés
    deputes = pd.DataFrame({
        "Parti": list(totaux_deputes_depuis_carte.keys()),
        "Sièges": list(totaux_deputes_depuis_carte.values())
    })
    
    # === 5. Appliquer le double seuil LOE (Art. 13-A)
    candidats_5pourcent = {p: v for p, v in pourcentages.items() if v >= SEUIL_ELECTORAL}
    deputes = deputes[deputes["Parti"].isin(candidats_5pourcent.keys())]
    deputes = deputes[deputes["Sièges"] >= 7]
    
    # === 6. Créer le dictionnaire filtré à utiliser pour toutes les autres étapes
    partis_admis = set(deputes["Parti"])
    pourcentages_filtrés = {p: v for p, v in pourcentages.items() if p in partis_admis}
    

    # === 6. Redistribuer les sièges perdus (pour obtenir exactement 130 sièges) ===

    total_initial = sum(totaux_deputes_depuis_carte.values())
    total_filtré = deputes["Sièges"].sum()
    sieges_a_redistribuer = total_initial - total_filtré

    if sieges_a_redistribuer > 0:
        # Proportion des partis admissibles
        total_votes_valides = sum([pourcentages[p] for p in partis_admis])
        proportions = {p: pourcentages[p] / total_votes_valides for p in partis_admis}

        # Attribution proportionnelle des sièges perdus
        for p in deputes["Parti"]:
            seats_to_add = round(proportions[p] * sieges_a_redistribuer)
            deputes.loc[deputes["Parti"] == p, "Sièges"] += seats_to_add

        # Correction pour atteindre exactement 130 (ajustement d'arrondi)
        difference = 130 - deputes["Sièges"].sum()
        if difference != 0:
            top_parti = deputes.sort_values("Sièges", ascending=False).iloc[0]["Parti"]
            deputes.loc[deputes["Parti"] == top_parti, "Sièges"] += difference



    # === 7. Régénérer les résultats régionaux uniquement avec les partis admis
    resultats_par_region = {}
    for district in DISTRIBUTION_DEPUTES:
        df_local = appliquer_cifra_par_district(
            pourcentages_filtrés,
            {district: DISTRIBUTION_DEPUTES[district]},
            perturbations=perturbations_par_district
        )
        resultats_par_region[district] = dict(zip(df_local["Parti"], df_local["Sièges"]))


   # senat, _ = simulateur_senat(pourcentages, pourcentages_filtrés)

    senat, region_senateurs = simulateur_senat(pourcentages, pourcentages_filtrés)
    senat["Sièges"] = senat["Sièges_majoritaire"] + senat["Sièges_Lima"] + senat["Sièges_National"]
    senat_affiche = senat[["Parti", "Sièges"]].copy()
    print(senat_affiche.to_string(index=False))

    senat["Sièges"] = senat["Sièges"].astype(int)
    generate_double_half_circle(deputes, senat)

    # Appliquer le seuil électoral national AVANT TOUT
    #pourcentages_filtrés = {p: v for p, v in pourcentages.items() if v >= SEUIL_ELECTORAL}
    
    print()
    print()
    print("============== Partis admis au Parlement (>= 5%) ==============")
    for parti, pct in sorted(pourcentages_filtrés.items(), key=lambda x: x[1], reverse=True):
        print(f"{parti}: {pct:.1f}%")
    
    print("\n============== Répartition des Députés ==============")
   # deputes = appliquer_cifra_par_district(pourcentages_filtrés, DISTRIBUTION_DEPUTES)
    print(deputes.to_string(index=False))

    print("\n============== Comparaison avec le Congrès de 2021 ==============")
    comparaison = comparer_resultats(deputes, resultats_2021)
    print(f"{'Parti':<25}{'2021':>6}{'Simulé':>8}{'Écart':>7}")
    for parti, s2021, s_sim, ecart in comparaison:
        print(f"{parti:<25}{s2021:>6}{s_sim:>8}{ecart:>7}")

    print("\n============== Répartition du Sénat ==============")
    senat_affiche = senat[["Parti", "Sièges"]].copy()
    print(senat_affiche.to_string(index=False))

    total_senat = senat["Total_Sénat"].sum()
    sieges_vacants = NB_SENATEURS_TOTAL - total_senat
    if sieges_vacants > 0:
        print(f"\n⚠️ {sieges_vacants} sièges vacants au Sénat en raison du seuil électoral de 5 %. Conformément à la LOE, ces sièges ne sont pas redistribués.")
    
    graphique_senat_proportionnel(senat)
    simuler_election_presidentielle(intentions_vote)
    generer_carte_senateurs_majoritaires(region_senateurs, r"C:\Matthieu\Documents\carte_senateurs.png")

  #  print(">>> Génération du graaphique en cours ...")

    stats_monte_carlo = simulation_monte_carlo(pourcentages)
    afficher_bornes_monte_carlo(stats_monte_carlo)

        # Exemple après ta simulation
  #  resultats = appliquer_cifra_par_district(pourcentages_filtrés, DISTRIBUTION_DEPUTES)
    
    # Transformer en dict de dicts région -> {parti: sièges}
    resultats_par_region = {}
    for district in DISTRIBUTION_DEPUTES:
        df_local = appliquer_cifra_par_district(
            pourcentages_filtrés,
            {district: DISTRIBUTION_DEPUTES[district]},
            perturbations=perturbations_par_district
        )
        resultats_par_region[district] = dict(zip(df_local["Parti"], df_local["Sièges"]))

        # Stockage console (facultatif, à adapter selon ce que tu veux capturer)
    
    buffer = StringIO()
    sys_stdout = sys.stdout
    sys.stdout = buffer
 
        # Générer la carte régionale à partir du parti dominant par district
    repartition_par_region = {}
    
    # Pour chaque district, on répète une mini-simulation et on prend le parti majoritaire
    for district, nb_sieges in DISTRIBUTION_DEPUTES.items():
        # Génère des votes simulés dans ce district
        perturbation = {p: max(0, v + random.uniform(-2, 2)) for p, v in pourcentages_filtrés.items()}
        total_valid = sum(perturbation.values())
        votes = {p: perturbation[p] / total_valid * 1000000 for p in perturbation}
        gagnant = max(votes.items(), key=lambda x: x[1])[0]
        repartition_par_region[district] = gagnant
    
  #  generer_carte_deputes_par_region(
  #      repartition_par_region,
  #      titre="Parti victorieux par départements (députés)",
  #      fichier_sortie=r"C:\Matthieu\Documents\carte_deputes.png"
   # )

    generer_carte_par_sieges(resultats_par_region, r"C:\Matthieu\Documents\carte_par_sieges.png")

    # Rejouer une mini synthèse textuelle
    print("=== Résumé de la simulation ===")
    for parti, pct in sorted(pourcentages_filtrés.items(), key=lambda x: x[1], reverse=True):
        print(f"{parti}: {pct:.1f}%")

    print("\nRépartition députés :")
    print(deputes.to_string(index=False))
    comparaison = comparer_resultats(deputes, resultats_2021)
    print(f"\n{'Parti':<25}{'2021':>6}{'Simulé':>8}{'Écart':>7}")
    for parti, s2021, s_sim, ecart in comparaison:
        print(f"{parti:<25}{s2021:>6}{s_sim:>8}{ecart:>7}")
    
    print("\nRépartition sénat :")
    print(senat.to_string(index=False))

    print("\nMonte Carlo :")
    print(stats_monte_carlo.to_string())

    sys.stdout = sys_stdout
    texte_console = buffer.getvalue()

    generer_rapport_word(
    chemin_dossier=r"C:\Matthieu\Documents",
    texte_console=texte_console,
    image_hemicycle_path=r"C:\Matthieu\Documents\double_hemicycle.png",
    image_montecarlo_path=r"C:\Matthieu\Documents\montecarlo_barres.png",
    image_carte_path=r"C:\Matthieu\Documents\carte_par_sieges.png",
    df_deputes=deputes,
    df_senat=senat,
    stats_monte_carlo=stats_monte_carlo
    )   
    enregistrer_resultats_excel(deputes, senat, stats_monte_carlo)

#################################################################################################################################
# SIMULATION GENERALE REPETITIVE
#################################################################################################################################

def simulation_globale_repetee(intentions_vote: Dict[str, float], n: int = 50):
    tous_deputes = []
    tous_senat = []

    for i in range(n):
        print(f"\n>>> Simulation {i+1}/{n} en cours...")
        pourcentages_filtrés = {p: v for p, v in intentions_vote.items() if v >= SEUIL_ELECTORAL}

        # Application pondérations
     #   deputes = appliquer_cifra_par_district(pourcentages_filtrés, DISTRIBUTION_DEPUTES)
       # deputes["Sièges"] = deputes.apply(
        #    lambda row: int(round(row["Sièges"] * pondérations_personnelles.get(row["Parti"], 1.0))), axis=1
       # )

        # Créer les résultats régionaux (comme pour la carte)
        # ➤ Transforme en dict de dicts pour la carte
        resultats_par_region = {}
        for district in DISTRIBUTION_DEPUTES:
            df_local = appliquer_cifra_par_district(pourcentages_filtrés, {district: DISTRIBUTION_DEPUTES[district]})
            resultats_par_region[district] = dict(zip(df_local["Parti"], df_local["Sièges"]))


        # ➤ Agréger tous les résultats régionaux
        totaux_deputes_depuis_carte = defaultdict(int)
        for region, resultats in resultats_par_region.items():
            for parti, nb_sieges in resultats.items():
                totaux_deputes_depuis_carte[parti] += nb_sieges
        
        # ➤ Convertir en DataFrame utilisable par le demi-hémicycle
        deputes = pd.DataFrame({
            "Parti": list(totaux_deputes_depuis_carte.keys()),
            "Sièges": list(totaux_deputes_depuis_carte.values())
        })

        senat, _ = simulateur_senat(intentions_vote, pourcentages_filtrés)

      #  senat["Total_Sénat"] = senat.apply(
        #    lambda row: int(round(row["Total_Sénat"] * pondérations_personnelles.get(row["Parti"], 1.0))), axis=1
       # )

        tous_deputes.append(deputes.set_index("Parti")["Sièges"])
        tous_senat.append(senat.set_index("Parti")["Total_Sénat"])

    # Calcul des moyennes
    moy_deputes = pd.concat(tous_deputes, axis=1).fillna(0).mean(axis=1).round(2).sort_values(ascending=False)
    moy_senat = pd.concat(tous_senat, axis=1).fillna(0).mean(axis=1).round(2).sort_values(ascending=False)

    print("\n============== Moyenne des résultats après", n, "simulations ==============\n")
    print("🟦 Députés (130 sièges) :")
    print(moy_deputes.to_string())

    print("\n🟥 Sénat (60 sièges) :")
    print(moy_senat.to_string())

    return moy_deputes, moy_senat

#################################################################################################################################
# SAISIR LES INTENTIONS DE VOTE REPETEES
#################################################################################################################################

if __name__ == "__main__":
    intentions_voteR = {
        "Fuerza Popular": 11.34,
        "Acción Popular": 9.02,
        "Perú Libre": 13.41,
        "Alianza para el Progreso": 7.54,
        "Renovación Popular": 9.33,
        "Avanza Pais": 7.54,
        "Juntos por el Perú": 6.59,
        "Somos Perú": 6.13,
        "Podemos Perú": 5.83,
        "Partido Morado": 5.42,
        "Victoria Nacional": 4.96,
        "Frepap": 4.58,
        "Unión por el Perú": 2.07,
        "PPC": 1.65,
        "PNP": 1.52,
        "FA": 1.05
    }
    simulation_globale_repetee(intentions_voteR, n=1)

#################################################################################################################################
# SAISIR LES INTENTIONS DE VOTE 
#################################################################################################################################
"""
if __name__ == "__main__":
    intentions_vote = {
        "Fuerza Popular": 11.34,
        "Acción Popular": 9.02,
        "Perú Libre": 13.41,
        "Alianza para el Progreso": 7.54,
        "Renovación Popular": 9.33,
        "Avanza Pais": 7.54,
        "Juntos por el Perú": 6.59,
        "Somos Perú": 6.13,
        "Podemos Perú": 5.83,
        "Partido Morado": 5.42,
        "Victoria Nacional": 4.96,
        "Frepap": 4.58,
        "Unión por el Perú": 2.07,
        "PPC": 1.65,
        "PNP": 1.52,
        "FA": 1.05
    }
    simulation_globale(intentions_vote)
"""
#################################################################################################################################
# SAISIR LES INTENTIONS DE VOTE 2
#################################################################################################################################

if __name__ == "__main__":
    intentions_vote = {
        "Fuerza Popular": 11.34,
        "Acción Popular": 9.02,
        "Perú Libre": 13.41,
        "Alianza para el Progreso": 7.54,
        "Renovación Popular": 9.33,
        "Avanza Pais": 7.54,
        "Juntos por el Perú": 6.59,
        "Somos Perú": 6.13,
        "Podemos Perú": 5.83,
        "Partido Morado": 5.42,
        "Victoria Nacional": 4.96,
        "Frepap": 4.58,
        "Unión por el Perú": 2.07,
        "PPC": 1.65,
        "PNP": 1.52,
        "FA": 1.05
    }
    simulation_globale(intentions_vote)
